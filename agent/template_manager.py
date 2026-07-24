"""Word 模板管理模块"""
import os
import shutil
import re
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class TemplateManager:
    """Word 模板管理器"""
    
    def __init__(self, template_dir: str = "./templates"):
        self.template_dir = template_dir
        os.makedirs(template_dir, exist_ok=True)
    
    def save_template(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """保存用户上传的模板"""
        # 验证文件格式
        if not filename.endswith('.docx'):
            return {"success": False, "error": "仅支持 .docx 格式"}
        
        # 生成唯一文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\-]', '_', filename)
        saved_path = os.path.join(self.template_dir, f"{timestamp}_{safe_name}")
        
        with open(saved_path, 'wb') as f:
            f.write(file_content)
        
        # 提取模板中的占位符
        placeholders = self._extract_placeholders(saved_path)
        
        return {
            "success": True,
            "template_id": timestamp,
            "filename": filename,
            "path": saved_path,
            "placeholders": placeholders
        }
    
    def _extract_placeholders(self, template_path: str) -> list:
        """提取模板中的占位符 {{变量名}}"""
        doc = Document(template_path)
        placeholders = set()
        
        # 检查段落
        for para in doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
            placeholders.update(matches)
        
        # 检查表格单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    placeholders.update(matches)
        
        return sorted(list(placeholders))
    
    def apply_template(self, template_path: str, content: str, variables: Dict[str, str]) -> str:
        """应用模板，替换占位符"""
        # 创建输出目录
        output_dir = "./output/documents"
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成输出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = os.path.basename(template_path).replace('.docx', '')
        output_path = os.path.join(output_dir, f"{base_name}_生成_{timestamp}.docx")
        
        # 复制模板
        shutil.copy(template_path, output_path)
        
        # 打开文档并替换占位符
        doc = Document(output_path)
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            if '{{' in para.text:
                replaced = False
                for key, value in variables.items():
                    if f'{{{{{key}}}}}' in para.text:
                        para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
                        replaced = True
                if replaced:
                    # 应用字体样式
                    for run in para.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{{' in cell.text:
                        for key, value in variables.items():
                            if f'{{{{{key}}}}}' in cell.text:
                                cell.text = cell.text.replace(f'{{{{{key}}}}}', str(value))
        
        # 添加生成信息
        doc.add_paragraph()
        info_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        for run in info_para.runs:
            run.font.name = "宋体"
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        
        doc.save(output_path)
        return output_path
    
    def list_templates(self) -> list:
        """列出所有可用模板"""
        templates = []
        for f in os.listdir(self.template_dir):
            if f.endswith('.docx'):
                path = os.path.join(self.template_dir, f)
                placeholders = self._extract_placeholders(path)
                templates.append({
                    "filename": f,
                    "path": path,
                    "placeholders": placeholders,
                    "size": os.path.getsize(path)
                })
        return templates
