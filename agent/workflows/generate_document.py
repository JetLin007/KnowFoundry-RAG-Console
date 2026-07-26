"""通用文档生成工作流 - 根据用户输入动态生成对应文档"""
from typing import TypedDict, List, Annotated, Dict, Any
import operator
import tempfile
import os
from datetime import datetime
import re

from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.prompts import ChatPromptTemplate

from qa_core.llm.client import get_chat_model
from agent.tools.rag_tool import RAGTool
from agent.workflows.base import BaseWorkflow
from agent.template_manager import TemplateManager

from docx import Document as DocxDocument
# ============================================
# GJB 438C 模板文档映射和结构提取
# ============================================

def extract_doc_structure_from_template(template_path: str) -> Dict[str, Any]:
    """从 Word 模板中提取文档结构（章节标题和层级）"""
    if not template_path or not os.path.exists(template_path):
        return {"structure": "未找到模板文档，请使用通用结构", "sections": []}
    
    try:
        doc = DocxDocument(template_path)
        sections = []
        current_level = 0
        structure_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测标题（通过样式或格式判断）
            # 简单判断：如果文本以数字开头或包含"章"、"节"等
            if para.style.name and 'Heading' in para.style.name:
                level = int(para.style.name.replace('Heading', '').strip()) if para.style.name.replace('Heading', '').strip().isdigit() else 1
                # 限制最大层级
                if level > 5:
                    level = 5
                sections.append({
                    "level": level,
                    "title": text,
                    "content": ""
                })
                # 生成结构描述
                indent = "  " * (level - 1)
                structure_lines.append(f"{indent}{text}")
            elif text.startswith('第') and ('章' in text or '节' in text):
                # 中文章节标题
                level = 1 if '章' in text else 2
                sections.append({
                    "level": level,
                    "title": text,
                    "content": ""
                })
                indent = "  " * (level - 1)
                structure_lines.append(f"{indent}{text}")
            elif text and len(text) < 50 and (text.isupper() or text[0].isdigit()):
                # 可能是标题
                sections.append({
                    "level": 1,
                    "title": text,
                    "content": ""
                })
                structure_lines.append(text)
        
        # 如果提取不到结构，使用默认结构
        if not sections:
            structure_lines = get_default_structure_for_doc_type(None)
            sections = []
        
        return {
            "structure": "\n".join(structure_lines),
            "sections": sections,
            "raw_text": "\n".join([p.text for p in doc.paragraphs[:50]])
        }
    except Exception as e:
        print(f"提取模板结构失败: {e}")
        return {"structure": "无法提取模板结构，请使用通用结构", "sections": []}

def get_default_structure_for_doc_type(doc_type: str) -> str:
    """获取文档类型的默认结构"""
    default_structures = {
        "开发计划": """1. 范围
   1.1 标识
   1.2 系统概述
   1.3 文档概述
   1.4 与其它计划的关系
2. 引用文档
3. 策划背景概述
4. 实施整个软件开发活动的计划
   4.1 软件开发过程
   4.2 软件开发总体计划
5. 实施详细软件开发活动的计划
   5.1 项目计划及监督
   5.2 建立软件开发环境
6. 进度安排及活动网络
7. 项目的组织和资源""",
        "质量保证": """1. 范围
   1.1 标识
   1.2 系统概述
   1.3 文档概述
   1.4 与其他计划之间的关系
2. 引用文档
3. 质量保证组织
4. 质量目标
5. 研制开发过程中的质量保证
6. 交付使用、售后服务的质量保证
7. 日程表
8. 质量审核和质量监督""",
        "技术方案": """1. 范围
   1.1 标识
   1.2 系统概述
   1.3 文档概述
2. 引用文档
3. 总体方案与结构
4. 模型设计方案
5. 资源配置方案
6. 系统配置方案
7. 接口配置方案
8. 组织机构及人员配置
9. 关键技术
10. 方案实施的技术路线和实施计划
11. 经费概算及规划""",
        "需求规格": """1. 范围
   1.1 标识
   1.2 系统概述
   1.3 文档概述
2. 引用文档
3. 需求
   3.1 要求的状态和方式
   3.2 软件能力需求
   3.3 软件外部接口需求
4. 合格性规定
5. 需求可追踪性""",
        "配置管理": """1. 范围
   1.1 标识
   1.2 系统概述
   1.3 文档概述
   1.4 与其他计划之间的关系
2. 引用文档
3. 组织和职责
4. 一般要求
5. 软件配置管理活动
6. 工具、技术和方法
7. 对供货单位的控制
8. 进度表""",
        "项目建设方案": """1. 项目概述
2. 需求分析
3. 总体建设方案
4. 软件设计方案
5. 项目实施计划
6. 运维保障方案
7. 投资概算"""
    }
    return default_structures.get(doc_type, default_structures.get("质量保证"))

# GJB 438C 模板文档路径映射
GJB_TEMPLATE_MAP = {
    "技术方案": {
        "doc_type": "技术方案",
        "template_path": "scenarios/military_software_438c/data/技术方案/[01]软件总体技术方案-438C.docx",
        "title": "软件总体技术方案",
        "gjb_section": "GJB 438C"
    },
    "开发计划": {
        "doc_type": "开发计划",
        "template_path": "scenarios/military_software_438c/data/开发计划/[02]软件开发计划-438C（共25页）.docx",
        "title": "软件开发计划",
        "gjb_section": "GJB 438C"
    },
    "配置管理": {
        "doc_type": "配置管理",
        "template_path": "scenarios/military_software_438c/data/配置管理/[03]软件配置管理计划-438C.docx",
        "title": "软件配置管理计划",
        "gjb_section": "GJB 438C"
    },
    "质量保证": {
        "doc_type": "质量保证",
        "template_path": "scenarios/military_software_438c/data/质量保证/[04]软件质量保证计划-438C.docx",
        "title": "软件质量保证计划",
        "gjb_section": "GJB 438C"
    },
    "需求规格": {
        "doc_type": "需求规格",
        "template_path": "scenarios/military_software_438c/data/需求规格/[08]软件需求规格说明-438C.docx",
        "title": "软件需求规格说明",
        "gjb_section": "GJB 438C"
    },
    "标准化": {
        "doc_type": "标准化",
        "template_path": "scenarios/military_software_438c/data/标准化/[05]软件标准化大纲-438C（21页）.docx",
        "title": "软件标准化大纲",
        "gjb_section": "GJB 438C"
    },
    "可靠性": {
        "doc_type": "可靠性",
        "template_path": "scenarios/military_software_438c/data/可靠性/[06]可靠性和可维护性大纲-438C.docx",
        "title": "软件可靠性和可维护性大纲",
        "gjb_section": "GJB 438C"
    },
    "安全性": {
        "doc_type": "安全性",
        "template_path": "scenarios/military_software_438c/data/安全性/[07]安全性大纲-438C.docx",
        "title": "软件安全性大纲",
        "gjb_section": "GJB 438C"
    },
    "审查报告": {
        "doc_type": "审查报告",
        "template_path": None,  # 使用通用模板
        "title": "技术审查报告",
        "gjb_section": "GJB 438C"
    },
    "用户手册": {
        "doc_type": "用户手册",
        "template_path": None,
        "title": "软件用户手册",
        "gjb_section": "GJB 438C"
    },
    "项目建设方案": {
        "doc_type": "项目建设方案",
        "template_path": "scenarios/enterprise_knowledge/data/项目建设方案.docx",  # 如果有模板
        "title": "项目建设方案",
        "gjb_section": "通用"
    }
}
# ============================================
# 文档状态定义
# ============================================

class DocumentState(TypedDict):
    user_request: str
    doc_type: str
    product_name: str
    style: str
    emphasis: str
    special_requirements: str
    format_settings: Dict[str, Any]
    retrieved_docs: List[Dict[str, Any]]
    doc_content: str
    doc_path: str
    messages: Annotated[List, operator.add]
    error: str


# ============================================
# 默认格式配置
# ============================================

DEFAULT_FORMAT = {
    "font_name": "宋体",
    "font_size": 12,
    "heading1_font": "黑体",
    "heading1_size": 18,
    "heading2_font": "黑体",
    "heading2_size": 16,
    "heading3_font": "黑体",
    "heading3_size": 14,
    "line_spacing": 1.5,
    "margin_top": 2.54,
    "margin_bottom": 2.54,
    "margin_left": 3.17,
    "margin_right": 3.17,
    "alignment": "left",
    "table_style": "Light Grid Accent 1",
    "page_orientation": "portrait"
}


# ============================================
# 文档类型映射
# ============================================

DOC_TYPE_MAP = {
    "质量保证": {
        "keywords": ["质量保证计划", "质量计划", "SQAP", "质量保证"],
        "title": "软件质量保证计划",
    },
    "开发计划": {
        "keywords": ["开发计划", "软件开发计划", "项目计划"],
        "title": "软件开发计划",
    },
    "配置管理": {
        "keywords": ["配置管理计划", "配置计划", "配置管理"],
        "title": "软件配置管理计划",
    },
    "需求规格": {
        "keywords": ["需求规格说明", "需求规格", "SRS", "需求", "功能需求"],
        "title": "软件需求规格说明",
    },
    "技术方案": {
        "keywords": ["总体技术方案", "技术方案", "架构设计"],
        "title": "软件总体技术方案",
    },
    "审查报告": {
        "keywords": ["审查报告", "技术审查", "评审报告"],
        "title": "技术审查报告",
    },
    "测试用例": {
        "keywords": ["测试用例", "测试计划", "测试"],
        "title": "软件测试用例文档",
    },
    "验收报告": {
        "keywords": ["验收报告", "验收", "交付验收"],
        "title": "软件验收报告",
    },
    "用户手册": {
        "keywords": ["用户手册", "使用手册", "操作手册"],
        "title": "软件用户手册",
    },
    "项目建设方案": {
        "keywords": ["项目建设方案", "建设方案", "项目方案", "实施方案"],
        "title": "项目建设方案",
        "structure": """1. 项目概述
   1.1 项目背景
   1.2 项目建设目标
   1.3 项目范围
   1.4 项目总投资估算
2. 需求分析
   2.1 业务需求
   2.2 用户需求
   2.3 功能需求
   2.4 非功能需求
3. 总体建设方案
   3.1 建设原则
   3.2 技术架构
   3.3 应用架构
   3.4 数据架构
   3.5 安全架构
4. 软件设计方案
   4.1 数据可视化系统设计
   4.2 三维电子沙盘软件设计
   4.3 各子系统功能模块
   4.4 系统集成方案
5. 数据可视化方案
   5.1 数据源接入
   5.2 数据清洗与处理
   5.3 可视化大屏设计
   5.4 图表类型与展示
6. 三维电子沙盘方案
   6.1 三维场景构建
   6.2 数据叠加展示
   6.3 交互操作设计
   6.4 性能优化策略
7. 项目实施计划
   7.1 项目阶段划分
   7.2 里程碑计划
   7.3 资源投入计划
   7.4 风险应对措施
8. 运维保障方案
   8.1 系统运维体系
   8.2 数据更新机制
   8.3 安全保障措施
9. 投资概算
   9.1 软硬件采购清单
   9.2 开发费用估算
   9.3 运维费用估算""",
        "prompt": """你是一位资深信息化项目咨询专家。根据提取的文档内容和用户要求，为 **{product_name}** 生成一份专业的项目建设方案。

## 文档结构要求

### 1. 项目概述
- 项目背景：阐述项目建设的原因和必要性
- 项目建设目标：明确项目要达到的具体目标
- 项目范围：界定项目建设的边界和范围
- 项目总投资估算：提供投资概算金额

### 2. 需求分析
- 业务需求：用户的业务需求描述
- 用户需求：各类用户的具体需求
- 功能需求：系统应具备的功能
- 非功能需求：性能、安全、可靠性等要求

### 3. 总体建设方案
- 建设原则：项目建设遵循的原则
- 技术架构：总体技术架构设计
- 应用架构：应用系统的组成和关系
- 数据架构：数据模型和数据流转
- 安全架构：安全保障体系设计

### 4. 软件设计方案
- 数据可视化系统设计：可视化系统架构和功能
- 三维电子沙盘软件设计：沙盘软件的设计方案
- 各子系统功能模块：各功能模块详细设计
- 系统集成方案：系统间的集成方式


### 5. 项目实施计划
- 项目阶段划分：项目实施的阶段划分
- 里程碑计划：关键里程碑节点
- 资源投入计划：人力、物力资源投入
- 风险应对措施：主要风险及应对策略

### 6. 运维保障方案
- 系统运维体系：运维组织和管理体系
- 数据更新机制：数据更新策略和频率
- 安全保障措施：系统安全保障方案

### 7. 投资概算
- 软硬件采购清单：需要采购的软硬件设备
- 开发费用估算：开发工作量及费用
- 运维费用估算：年度运维费用

## 格式要求
- 使用标准章节编号（1. 2. 3. ...）
- 关键内容使用表格呈现
- 重要数据使用加粗标注
- 专业术语需定义说明
- 内容应具备可操作性"""
    }
}


# ============================================
# 风格配置
# ============================================

STYLE_CONFIG = {
    "简略": {
        "max_sections": 3,
        "detail_level": "概述性",
        "prompt_suffix": "生成简洁的概述性文档，每个部分不超过3点，总字数控制在500字左右。"
    },
    "标准": {
        "max_sections": 5,
        "detail_level": "标准",
        "prompt_suffix": "生成标准格式的文档，每个部分覆盖主要内容，总字数控制在1500字左右。"
    },
    "详细": {
        "max_sections": 8,
        "detail_level": "详尽",
        "prompt_suffix": "生成详细的完整文档，每个部分展开论述，包含具体措施和量化指标，总字数控制在3000字以上。"
    }
}

EMPHASIS_CONFIG = {
    "风险": {
        "prompt_suffix": "重点关注风险识别、风险评估、风险应对和风险监控，在每个部分突出风险相关内容。"
    },
    "质量": {
        "prompt_suffix": "重点关注质量标准、质量度量、质量审核和质量改进，在每个部分突出质量相关内容。"
    },
    "进度": {
        "prompt_suffix": "重点关注进度安排、里程碑、资源调配和进度监控，在每个部分突出进度相关内容。"
    },
    "安全": {
        "prompt_suffix": "重点关注安全性设计、安全测试、安全评估和安全保障，在每个部分突出安全相关内容。"
    },
    "成本": {
        "prompt_suffix": "重点关注成本估算、预算控制、成本效益分析和成本优化，在每个部分突出成本相关内容。"
    }
}


# ============================================
# 工作流主类
# ============================================

class GenerateDocumentWorkflow(BaseWorkflow):
    """通用文档生成工作流"""

    def __init__(self, scenario_id: str = "enterprise_knowledge"):
        self.rag_tool = RAGTool(scenario_id=scenario_id)
        self.llm = get_chat_model(streaming=False)
        self.graph = self._build_graph()
        self.memory = MemorySaver()
        self.app = self.graph.compile(checkpointer=self.memory)

    def _extract_doc_type(self, request: str) -> str:
        for doc_type, config in DOC_TYPE_MAP.items():
            for keyword in config["keywords"]:
                if keyword in request:
                    return doc_type
        return "质量保证"

    def _extract_product_name(self, request: str) -> str:
        patterns = [
            r'[为给]?\s*([^\s，,。的]{2,20})\s*产品',
            r'产品\s*([^\s，,。]{2,20})',
            r'项目\s*([^\s，,。]{2,20})',
            r'[为给]\s*([^\s，,。]{2,20})\s*[生制]',
        ]
        for pattern in patterns:
            match = re.search(pattern, request)
            if match:
                return match.group(1).strip()
        return "指定产品"

    def _extract_style(self, request: str) -> str:
        if "简略" in request or "简要" in request or "概述" in request:
            return "简略"
        elif "详细" in request or "完整" in request or "详尽" in request:
            return "详细"
        return "标准"

    def _extract_emphasis(self, request: str) -> str:
        if "风险" in request or "风险控制" in request:
            return "风险"
        elif "质量" in request or "质量控制" in request:
            return "质量"
        elif "进度" in request or "时间" in request or "里程碑" in request:
            return "进度"
        elif "安全" in request or "安全性" in request:
            return "安全"
        elif "成本" in request or "预算" in request:
            return "成本"
        return "质量"

    def _extract_special_requirements(self, request: str) -> str:
        special_terms = [
            "符合GJB", "符合标准", "满足规范",
            "强调", "突出", "重点关注",
            "包含", "涵盖", "包括",
            "适用于", "针对"
        ]
        requirements = []
        for term in special_terms:
            if term in request:
                requirements.append(term)
        return "、".join(requirements) if requirements else "无特殊要求"

    def _extract_format_settings(self, request: str) -> Dict[str, Any]:
        """从用户输入中提取格式设置（增强版）"""
        format_settings = DEFAULT_FORMAT.copy()
        
        # 字体映射
        font_map = {
            "宋体": "宋体",
            "黑体": "黑体", 
            "仿宋": "仿宋",
            "楷体": "楷体",
            "微软雅黑": "微软雅黑",
            "Times New Roman": "Times New Roman",
            "Arial": "Arial",
            "Calibri": "Calibri"
        }
        
        # 提取正文字体
        for key, value in font_map.items():
            if key in request:
                format_settings["font_name"] = value
                break
        
        # 提取标题字体（如果有"黑体标题"或"标题黑体"等关键词）
        if "黑体标题" in request or "标题黑体" in request:
            format_settings["heading1_font"] = "黑体"
            format_settings["heading2_font"] = "黑体"
            format_settings["heading3_font"] = "黑体"
        elif "宋体标题" in request or "标题宋体" in request:
            format_settings["heading1_font"] = "宋体"
            format_settings["heading2_font"] = "宋体"
            format_settings["heading3_font"] = "宋体"
        
        # 提取字号
        size_patterns = [
            r'(小?[三四]?号|初号|小初|大?一?二?三?四?五?号)',
            r'(\d{1,2})pt'
        ]
        for pattern in size_patterns:
            match = re.search(pattern, request)
            if match:
                try:
                    size = int(re.search(r'\d+', match.group()).group())
                    if 8 <= size <= 72:
                        format_settings["font_size"] = size
                        format_settings["heading1_size"] = min(size + 6, 28)
                        format_settings["heading2_size"] = min(size + 4, 24)
                        format_settings["heading3_size"] = min(size + 2, 20)
                except:
                    pass
                break
        
        # 提取对齐方式
        if "居中" in request:
            format_settings["alignment"] = "center"
        elif "右对齐" in request:
            format_settings["alignment"] = "right"
        elif "两端对齐" in request:
            format_settings["alignment"] = "justify"
        else:
            format_settings["alignment"] = "left"
        
        # 提取行距
        if "1.5倍" in request or "一倍半" in request:
            format_settings["line_spacing"] = 1.5
        elif "双倍" in request or "2倍" in request:
            format_settings["line_spacing"] = 2
        elif "单倍" in request:
            format_settings["line_spacing"] = 1
        
        # 提取页面方向
        if "横向" in request or "横版" in request:
            format_settings["page_orientation"] = "landscape"
        
        return format_settings

    def _get_doc_config(self, doc_type: str) -> Dict[str, Any]:
        return DOC_TYPE_MAP.get(doc_type, DOC_TYPE_MAP["质量保证"])

    def _build_graph(self):
        graph = StateGraph(DocumentState)
        graph.add_node("extract_info", self.extract_info)
        graph.add_node("retrieve_docs", self.retrieve_docs)
        graph.add_node("generate_document", self.generate_document)
        graph.add_node("save_document", self.save_document)
        graph.set_entry_point("extract_info")
        graph.add_edge("extract_info", "retrieve_docs")
        graph.add_edge("retrieve_docs", "generate_document")
        graph.add_edge("generate_document", "save_document")
        graph.add_edge("save_document", END)
        return graph

    def extract_info(self, state: DocumentState) -> Dict[str, Any]:
        request = state["user_request"]
        return {
            "doc_type": self._extract_doc_type(request),
            "product_name": self._extract_product_name(request),
            "style": self._extract_style(request),
            "emphasis": self._extract_emphasis(request),
            "special_requirements": self._extract_special_requirements(request),
            "format_settings": self._extract_format_settings(request)
        }

    def retrieve_docs(self, state: DocumentState) -> Dict[str, Any]:
        config = self._get_doc_config(state["doc_type"])
        query = f"{config['title']} {state['product_name']} {state['emphasis']}"
        docs = self.rag_tool.retrieve(query, top_k=5)
        mapped_docs = []
        for doc in docs:
            if 'text' in doc:
                mapped_docs.append({'content': doc.get('text', ''), 'score': doc.get('score', 0)})
            else:
                mapped_docs.append(doc)
        return {"retrieved_docs": mapped_docs}

    def _build_custom_prompt(self, state: DocumentState) -> str:
        """根据提取的信息构建自定义提示词，优先使用 DOC_TYPE_MAP 中的结构"""
        config = self._get_doc_config(state["doc_type"])
        style_info = STYLE_CONFIG.get(state["style"], STYLE_CONFIG["标准"])
        emphasis_info = EMPHASIS_CONFIG.get(state["emphasis"], EMPHASIS_CONFIG["质量"])
        fmt = state.get("format_settings", DEFAULT_FORMAT)
        product_name = state.get('product_name', '指定产品')
        doc_type = state.get('doc_type', '质量保证')
        
        # ===== 1. 优先使用 DOC_TYPE_MAP 中的 structure 和 prompt =====
        structure = config.get("structure", "")
        custom_prompt = config.get("prompt", "")
        
        # 如果没有 structure，尝试从 GJB 模板提取
        if not structure:
            template_info = GJB_TEMPLATE_MAP.get(doc_type)
            if template_info:
                template_path = template_info.get("template_path")
                if template_path and os.path.exists(template_path):
                    print(f"📄 从 GJB 438C 模板提取结构: {template_path}")
                    template_data = extract_doc_structure_from_template(template_path)
                    structure = template_data.get("structure")
                    if structure and "未提取" not in structure:
                        print(f"✅ 成功提取模板结构")
            
            # 如果还是没有，使用默认结构
            if not structure:
                structure = get_default_structure_for_doc_type(doc_type)
                print(f"⚠️ 使用默认结构: {doc_type}")
        
        # ===== 2. 获取文档标题 =====
        doc_title = config.get('title', doc_type)
        
        # ===== 3. 构建最终提示词 =====
        if custom_prompt:
            try:
                prompt = custom_prompt.format(product_name=product_name)
            except KeyError:
                prompt = custom_prompt
            prompt += f"\n\n### 文档结构要求\n{structure}"
            prompt += f"\n\n### 格式要求\n- 正文字体：{fmt.get('font_name', '宋体')}，字号：{fmt.get('font_size', 12)}pt\n- 标题字体：{fmt.get('heading1_font', '黑体')}\n- 行距：{fmt.get('line_spacing', 1.5)}倍\n- 对齐方式：{fmt.get('alignment', '左对齐')}"
        else:
            # 判断是否使用 GJB 438C 模板
            template_info = GJB_TEMPLATE_MAP.get(doc_type)
            gjb_ref = f"GJB 438C-2021《军用软件开发文档通用要求》" if template_info else ""
            
            prompt = f"""你是一位军用软件开发文档专家。为产品 **{product_name}** 生成一份 **{doc_title}**。

    ## 文档要求
    {f"- 严格遵循 {gjb_ref}" if gjb_ref else "- 遵循标准文档格式"}
    - 文档结构必须包含以下章节

    ## 文档结构
    {structure}

    ## 内容要求
    - 每个章节必须包含实质性内容
    - 内容应与产品 **{product_name}** 的具体情况相关
    - 专业术语需定义说明
    - 关键信息应使用表格呈现

    **文档风格要求**：{style_info['prompt_suffix']}

    **侧重要求**：{emphasis_info['prompt_suffix']}

    **特殊要求**：{state['special_requirements']}

    **格式要求**：
    - 正文字体：{fmt.get('font_name', '宋体')}，字号：{fmt.get('font_size', 12)}pt
    - 标题字体：{fmt.get('heading1_font', '黑体')}
    - 行距：{fmt.get('line_spacing', 1.5)}倍
    - 对齐方式：{fmt.get('alignment', '左对齐')}

    请确保文档内容专业、完整{f"，符合GJB 438C标准要求" if gjb_ref else ""}。
    """
        
        return prompt

    def generate_document(self, state: DocumentState) -> Dict[str, Any]:
        config = self._get_doc_config(state["doc_type"])
        docs = state.get("retrieved_docs", [])
        
        if not docs:
            return {"doc_content": f"未能检索到相关文档，请确保知识库中包含{config['title']}相关内容。"}
        
        docs_text = "\n".join([doc.get('content', '') for doc in docs])
        system_prompt = self._build_custom_prompt(state)
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "参考文档内容：\n{docs}\n\n请生成文档：")
        ])
        
        response = self.llm.invoke(prompt.format_messages(docs=docs_text))
        return {"doc_content": response.content}

    def save_document(self, state: DocumentState) -> Dict[str, Any]:
        """保存文档为 Word（GJB 438C 格式）"""
        config = self._get_doc_config(state["doc_type"])
        fmt = state.get("format_settings", DEFAULT_FORMAT)
        product_name = state.get('product_name', '指定产品')
        doc_title = config.get('title', '文档')
        
        output_dir = "./output/documents"
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        product = product_name.replace(" ", "_")
        style = state["style"]
        title = doc_title.replace(" ", "_")
        file_path = os.path.join(output_dir, f"{product}_{title}_{style}_{timestamp}.docx")
        
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import re
            
            content = state["doc_content"]
            doc = Document()
            
            # ==========================================
            # 1. 设置页面格式（A4，符合GJB 438C）
            # ==========================================
            section = doc.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            
            # ==========================================
            # 2. 封面页
            # ==========================================
            # 密级（右上角）
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run("密级：内部")
            run.font.name = "黑体"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
            
            # 空行
            doc.add_paragraph()
            
            # 文档类型（居中，大号）
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(doc_title)
            run.font.name = "黑体"
            run.font.size = Pt(22)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
            
            # 产品型号
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("产品型号-XXXX")
            run.font.name = "宋体"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            
            # 空行
            for _ in range(3):
                doc.add_paragraph()
            
            # 编制单位
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("XXXXXXXX公司")
            run.font.name = "宋体"
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            
            # 日期
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{datetime.now().strftime('%Y年%m月%d日')}")
            run.font.name = "宋体"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            
            # 分页
            doc.add_page_break()
            
            # ==========================================
            # 3. 签署页
            # ==========================================
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(doc_title)
            run.font.name = "黑体"
            run.font.size = Pt(20)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("产品型号-XXXX")
            run.font.name = "宋体"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            
            # 空行
            for _ in range(4):
                doc.add_paragraph()
            
            # 签署栏（左右两列）
            signatures = [
                ("编制：", "日期："),
                ("校对：", "日期："),
                ("审核：", "日期："),
                ("标审：", "日期："),
                ("审定：", "日期："),
                ("批准：", "日期：")
            ]
            
            for signer, date in signatures:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"{signer}        {date}")
                run.font.name = "宋体"
                run.font.size = Pt(14)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            
            # 分页
            doc.add_page_break()
            
            # ==========================================
            # 4. 正文内容
            # ==========================================
            # 设置正文样式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            style.paragraph_format.line_spacing = Pt(20)  # 固定行距20磅
            style.paragraph_format.first_line_indent = Cm(0.75)  # 首行缩进2字符
            
            # 解析内容
            lines = content.split('\n')
            in_table = False
            table_data = []
            
            # 定义标题样式
            heading_styles = {
                'h1': {'font': '黑体', 'size': 18, 'indent': 0, 'space_before': 12, 'space_after': 6},
                'h2': {'font': '黑体', 'size': 16, 'indent': 0, 'space_before': 10, 'space_after': 4},
                'h3': {'font': '黑体', 'size': 14, 'indent': 0, 'space_before': 8, 'space_after': 3},
                'h4': {'font': '黑体', 'size': 12, 'indent': 0, 'space_before': 6, 'space_after': 2},
            }
            
            for line in lines:
                line = line.rstrip()
                if not line:
                    continue
                
                # 检测标题
                header_level = None
                if line.startswith('# '):
                    header_level = 'h1'
                    text = line[2:].strip()
                elif line.startswith('## '):
                    header_level = 'h2'
                    text = line[3:].strip()
                elif line.startswith('### '):
                    header_level = 'h3'
                    text = line[4:].strip()
                elif line.startswith('#### '):
                    header_level = 'h4'
                    text = line[5:].strip()
                else:
                    text = line
                
                if header_level:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.space_before = Pt(heading_styles[header_level]['space_before'])
                    p.paragraph_format.space_after = Pt(heading_styles[header_level]['space_after'])
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(20)
                    
                    run = p.add_run(text)
                    run.font.name = heading_styles[header_level]['font']
                    run.font.size = Pt(heading_styles[header_level]['size'])
                    run.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), heading_styles[header_level]['font'])
                else:
                    # 普通段落
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0.75)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(20)
                    
                    for run in p.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            
            # ==========================================
            # 5. 生成目录（在正文前插入）
            # ==========================================
            # 在正文前插入目录
            # 由于python-docx不支持自动生成目录，我们创建手动目录
            # 从内容中提取标题
            
            # 重新解析内容，提取所有标题用于目录
            toc_entries = []
            for line in content.split('\n'):
                line = line.rstrip()
                if line.startswith('# '):
                    level = 1
                    text = line[2:].strip()
                elif line.startswith('## '):
                    level = 2
                    text = line[3:].strip()
                elif line.startswith('### '):
                    level = 3
                    text = line[4:].strip()
                elif line.startswith('#### '):
                    level = 4
                    text = line[5:].strip()
                else:
                    continue
                
                if text and len(text) < 50:
                    toc_entries.append((level, text))
            
            # 在正文前插入目录页
            # 由于我们已经添加了封面和签署页，现在在正文前插入目录
            # 使用 python-docx 的目录功能（简化版）
            # 注意：实际生成需要更复杂的处理
            
            # 简单目录（手动创建）
            # 由于python-docx不支持自动TOC，我们用段落模拟
            # 在文档开头插入目录标题
            # 实际上，更合适的方式是生成后手动更新目录
            
            # 在分页之前，回到文档开头插入目录
            # 我们简单处理：在正文前添加目录标题和条目
            
            # 由于复杂的目录生成需要更多代码，这里简化处理
            # 实际项目中可以使用 python-docx 的 add_toc 功能或第三方库
            
            doc.save(file_path)
            return {"doc_path": file_path, "format_applied": fmt}
            
        except ImportError:
            md_path = file_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(state["doc_content"])
            return {"doc_path": md_path, "warning": "python-docx 未安装，已保存为 Markdown 格式"}
        except Exception as e:
            md_path = file_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(state["doc_content"])
            return {"doc_path": md_path, "warning": f"Word 生成失败，已保存为 Markdown: {str(e)}"}

    async def run(self, request: str, thread_id: str = "default") -> Dict[str, Any]:
        config = {"configurable": {"thread_id": thread_id}}
        initial_state = {"user_request": request}
        final_state = await self.app.ainvoke(initial_state, config)

        return {
            "content": final_state.get("doc_content", ""),
            "doc_path": final_state.get("doc_path", ""),
            "doc_type": final_state.get("doc_type", "未知类型"),
            "product_name": final_state.get("product_name", "未知产品"),
            "style": final_state.get("style", "标准"),
            "emphasis": final_state.get("emphasis", "质量"),
            "format_settings": final_state.get("format_settings", DEFAULT_FORMAT)
        }
