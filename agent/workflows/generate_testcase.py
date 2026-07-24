"""生成测试用例工作流 - 第一个 LangGraph 实现"""
from typing import TypedDict, List, Annotated, Dict, Any
import operator
import tempfile
import os

from langgraph.graph import StateGraph, END
#from langgraph.checkpoint import MemorySaver
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage

from qa_core.llm.client import get_chat_model
from agent.tools.rag_tool import RAGTool
from agent.workflows.base import BaseWorkflow

class TestCaseState(TypedDict):
    """工作流状态"""
    user_request: str
    retrieved_docs: List[Dict[str, Any]]
    test_items: List[str]
    testcase_content: str
    doc_path: str
    messages: Annotated[List, operator.add]
    error: str

class GenerateTestcaseWorkflow(BaseWorkflow):
    """生成测试用例工作流"""
    def __init__(self, scenario_id: str = "enterprise_knowledge"):
        self.rag_tool = RAGTool(scenario_id=scenario_id)
        self.llm = get_chat_model(streaming=False)
        self.graph = self._build_graph()
        self.memory = MemorySaver()
        self.app = self.graph.compile(checkpointer=self.memory)     

    def _build_graph(self):
        graph = StateGraph(TestCaseState)

        graph.add_node("retrieve_docs", self.retrieve_docs)
        graph.add_node("extract_test_items", self.extract_test_items)
        graph.add_node("generate_testcase", self.generate_testcase)
        graph.add_node("save_document", self.save_document)

        graph.set_entry_point("retrieve_docs")
        graph.add_edge("retrieve_docs", "extract_test_items")
        graph.add_edge("extract_test_items", "generate_testcase")
        graph.add_edge("generate_testcase", "save_document")
        graph.add_edge("save_document", END)

        return graph

    def retrieve_docs(self, state: TestCaseState) -> Dict[str, Any]:
        """检索相关文档"""
        print("=" * 50)
        print("DEBUG: retrieve_docs 开始")
        print(f"DEBUG: 查询: {state.get('user_request', '')}")
        print(f"DEBUG: rag_tool 类型: {type(self.rag_tool)}")
        print(f"DEBUG: rag_tool scenario_id: {self.rag_tool.scenario_id}")
        print(f"DEBUG: rag_tool collection_name: {self.rag_tool.collection_name}")
        
        # 尝试不同的查询词
        queries = [
            state.get('user_request', ''),
            "触控校准 手势识别 红外感应",
            "红外触控嵌入式软件需求规范"
        ]
        
        for i, q in enumerate(queries):
            docs = self.rag_tool.retrieve(q, top_k=3)
            print(f"DEBUG: 查询{i+1} '{q[:30]}...' 返回 {len(docs)} 条")
            if docs:
                print(f"DEBUG: 第一条: {docs[0].get('text', '')[:100]}...")
                break
        
        # 使用最后一个查询的结果
        docs = self.rag_tool.retrieve(queries[-1], top_k=3)
        print(f"DEBUG: 最终使用查询 '{queries[-1][:30]}...' 返回 {len(docs)} 条")
        
        # 映射
        mapped_docs = []
        for doc in docs:
            if 'text' in doc:
                mapped_docs.append({
                    'content': doc.get('text', ''),
                    'score': doc.get('score', 0)
                })
            elif 'content' in doc:
                mapped_docs.append(doc)
            else:
                mapped_docs.append({'content': str(doc), 'score': 0})
        
        print(f"DEBUG: 映射后 {len(mapped_docs)} 条文档")
        result = {"retrieved_docs": mapped_docs}
        print("=" * 50)
        return result

#     def retrieve_docs(self, state: TestCaseState) -> Dict[str, Any]:
#         """直接注入模拟文档内容（跳过 RAG 检索）"""
#         mock_content = """
#     # 红外触控嵌入式软件需求规范
# 
#     ## 1. 功能需求
#     ### 1.1 触控校准
#     - 系统启动后自动进入触控校准模式
#     - 用户点击屏幕上的5个校准点完成校准
#     - 校准数据保存到非易失性存储器
# 
#     ### 1.2 手势识别
#     - 支持单指点击、双指缩放、三指滑动
#     - 手势识别响应时间 < 50ms
#     - 误触率 < 1%
# 
#     ### 1.3 红外感应
#     - 感应距离：0-50cm
#     - 感应角度：±30度
#     - 采样率：100Hz
# 
#     ## 2. 性能要求
#     - 系统启动时间 < 3秒
#     - CPU使用率 < 60%
#     - 内存占用 < 128MB
# 
#     ## 3. 接口规范
#     - 使用 I2C 协议与主控通信
#     - 通信速率：400kHz
#     - 数据包格式：8字节数据帧
# 
#     ## 4. 可靠性要求
#     - 连续运行 72 小时无故障
#     - 工作温度：-20℃ ~ 70℃
#     - 静电防护等级：±4kV
#     """
#         docs = [{"content": mock_content}]
#         return {"retrieved_docs": docs}

    def extract_test_items(self, state: TestCaseState) -> Dict[str, Any]:
        """从文档中提取测试项"""
        print("=" * 50)
        print("DEBUG: extract_test_items 开始")
        print(f"DEBUG: state keys: {state.keys()}")
        print(f"DEBUG: retrieved_docs 在 state 中: {'retrieved_docs' in state}")
        if 'retrieved_docs' in state:
            print(f"DEBUG: retrieved_docs 类型: {type(state['retrieved_docs'])}")
            print(f"DEBUG: retrieved_docs 长度: {len(state['retrieved_docs'])}")
            if state['retrieved_docs']:
                print(f"DEBUG: 第一条文档内容: {state['retrieved_docs'][0].get('content', '')[:100]}...")
        
        docs = state.get("retrieved_docs", [])
        print(f"DEBUG: 获取到 {len(docs)} 条文档")
        
        if not docs:
            return {"test_items": ["文档内容为空，无法提取测试项。"]}
        
        docs_text = "\n".join([doc.get('content', '') for doc in docs])
        print(f"DEBUG: 合并后文档长度: {len(docs_text)}")
        print("=" * 50)
        
        if not docs_text.strip():
            return {"test_items": ["文档内容为空，无法提取测试项。"]}
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", "你是一位软件测试专家。从以下软件开发规范/需求文档中提取关键的测试项或测试指标，以列表形式返回。"),
            ("human", "文档内容：\n{docs}\n\n请提取测试项：")
        ])
        response = self.llm.invoke(prompt.format_messages(docs=docs_text))
        items = [line.strip("- ").strip() for line in response.content.split("\n") if line.strip()]
        return {"test_items": items if items else ["无法提取测试项，请检查文档内容。"]}

    def generate_testcase(self, state: TestCaseState) -> Dict[str, Any]:
        """生成软件测试用例文档"""
        items = state.get("test_items", [])
        print(f"DEBUG: generate_testcase 收到 {len(items)} 个测试项")
        
        if not items or items == ["文档内容为空，无法提取测试项。"]:
            return {"testcase_content": "未能提取到测试项，请确保需求文档包含可测试的功能描述。"}
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", "你是一位资深软件测试工程师。根据提取的测试项和规范，生成一份详细的软件测试用例文档（Markdown格式）。"),
            ("human", "测试项：\n{items}\n\n请生成软件测试用例：")
        ])
        items_text = "\n".join(items)
        response = self.llm.invoke(prompt.format_messages(items=items_text))
        return {"testcase_content": response.content}

    def save_document(self, state: TestCaseState) -> Dict[str, Any]:
        """将 Markdown 转换为 Word 文档"""
        import tempfile
        import os
        from datetime import datetime
    
        # 创建输出目录
        output_dir = "./output/testcases"
        os.makedirs(output_dir, exist_ok=True)
    
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(output_dir, f"testcase_{timestamp}.docx")
    
        try:
            from docx import Document
            content = state["testcase_content"]
            doc = Document()
            doc.add_heading("软件测试用例文档", 0)
        
            for line in content.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], 1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], 3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
        
            doc.save(file_path)
            return {"doc_path": file_path}
        except ImportError:
            # 降级保存为 Markdown
            md_path = file_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(state["testcase_content"])
            return {"doc_path": md_path, "warning": "python-docx not installed, saved as Markdown"}

    async def run(self, request: str, thread_id: str = "default") -> Dict[str, Any]:
        """运行工作流"""
        config = {"configurable": {"thread_id": thread_id}}
        initial_state = {"user_request": request}
        final_state = await self.app.ainvoke(initial_state, config)
        return {
            "content": final_state.get("testcase_content", ""),
            "doc_path": final_state.get("doc_path", ""),
            "test_items": final_state.get("test_items", [])
        }
