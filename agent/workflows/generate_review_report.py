"""生成军工技术审查报告工作流"""
from typing import TypedDict, List, Annotated, Dict, Any
import operator
import tempfile
import os
from datetime import datetime

from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.prompts import ChatPromptTemplate

from qa_core.llm.client import get_chat_model
from agent.tools.rag_tool import RAGTool
from agent.workflows.base import BaseWorkflow

class ReviewReportState(TypedDict):
    user_request: str
    retrieved_docs: List[Dict[str, Any]]
    review_items: List[str]
    report_content: str
    doc_path: str
    messages: Annotated[List, operator.add]
    error: str

class GenerateReviewReportWorkflow(BaseWorkflow):
    """生成技术审查报告工作流"""

    def __init__(self, scenario_id: str = "defense_review_qa"):
        self.rag_tool = RAGTool(scenario_id=scenario_id)
        self.llm = get_chat_model(streaming=False)
        self.graph = self._build_graph()
        self.memory = MemorySaver()
        self.app = self.graph.compile(checkpointer=self.memory)

    def _build_graph(self):
        graph = StateGraph(ReviewReportState)

        graph.add_node("retrieve_docs", self.retrieve_docs)
        graph.add_node("extract_review_items", self.extract_review_items)
        graph.add_node("generate_report", self.generate_report)
        graph.add_node("save_document", self.save_document)

        graph.set_entry_point("retrieve_docs")
        graph.add_edge("retrieve_docs", "extract_review_items")
        graph.add_edge("extract_review_items", "generate_report")
        graph.add_edge("generate_report", "save_document")
        graph.add_edge("save_document", END)

        return graph

    def retrieve_docs(self, state: ReviewReportState) -> Dict[str, Any]:
        """检索相关文档"""
        docs = self.rag_tool.retrieve(state["user_request"], top_k=5)
        mapped_docs = []
        for doc in docs:
            if 'text' in doc:
                mapped_docs.append({'content': doc.get('text', ''), 'score': doc.get('score', 0)})
            else:
                mapped_docs.append(doc)
        return {"retrieved_docs": mapped_docs}

    def extract_review_items(self, state: ReviewReportState) -> Dict[str, Any]:
        """从文档中提取审查要点"""
        docs = state.get("retrieved_docs", [])
        if not docs:
            return {"review_items": ["未找到相关审查文档。"]}
        
        docs_text = "\n".join([doc.get('content', '') for doc in docs])
        prompt = ChatPromptTemplate.from_messages([
            ("system", """你是一位军工产品技术审查专家。从以下文档中提取技术审查的要点和检查项，以列表形式返回。

提取原则：
- 技术审查输入文件要求
- 评审会议组织要点
- 技术状态管理要求
- 试验验证要求
- 质量保证要点
- 风险识别与控制
"""),
            ("human", "文档内容：\n{docs}\n\n请提取审查要点：")
        ])
        response = self.llm.invoke(prompt.format_messages(docs=docs_text))
        items = [line.strip("- ").strip() for line in response.content.split("\n") if line.strip()]
        return {"review_items": items if items else ["无法提取审查要点。"]}

    def generate_report(self, state: ReviewReportState) -> Dict[str, Any]:
        """生成技术审查报告"""
        items = state.get("review_items", [])
        if not items:
            return {"report_content": "未能提取到审查要点，请确保文档包含技术审查相关内容。"}
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", """你是一位军工产品质量工程师。根据提取的审查要点，生成一份技术审查报告（Markdown格式）。

报告结构：
1. 审查概述
2. 审查依据（引用相关标准/文件）
3. 审查项目及结果
4. 发现的问题及建议
5. 审查结论
6. 后续行动计划
"""),
            ("human", "审查要点：\n{items}\n\n请生成技术审查报告：")
        ])
        items_text = "\n".join(items)
        response = self.llm.invoke(prompt.format_messages(items=items_text))
        return {"report_content": response.content}

    def save_document(self, state: ReviewReportState) -> Dict[str, Any]:
        """保存报告为 Word 文档"""
        output_dir = "./output/reports"
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(output_dir, f"technical_review_report_{timestamp}.docx")
        
        try:
            from docx import Document
            content = state["report_content"]
            doc = Document()
            doc.add_heading("技术审查报告", 0)
            
            for line in content.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], 1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], 3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
            
            doc.save(file_path)
            return {"doc_path": file_path}
        except ImportError:
            md_path = file_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(state["report_content"])
            return {"doc_path": md_path, "warning": "已保存为 Markdown 格式"}

    async def run(self, request: str, thread_id: str = "default") -> Dict[str, Any]:
        config = {"configurable": {"thread_id": thread_id}}
        initial_state = {"user_request": request}
        final_state = await self.app.ainvoke(initial_state, config)
        return {
            "content": final_state.get("report_content", ""),
            "doc_path": final_state.get("doc_path", ""),
            "review_items": final_state.get("review_items", [])
        }
